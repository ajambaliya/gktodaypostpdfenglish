import io
import os
import requests
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime
import pymongo
import asyncio
import telegram
import tempfile
import subprocess
import logging
from logging.handlers import RotatingFileHandler

# Setup logging
LOG_FILE = 'script_debug.log'
logger = logging.getLogger('ScriptLogger')
logger.setLevel(logging.DEBUG)

# File handler that writes log messages to a file, and rotates the log every 10MB, keeping 5 backups
file_handler = RotatingFileHandler(LOG_FILE, maxBytes=10 * 1024 * 1024, backupCount=5)
file_handler.setLevel(logging.DEBUG)

# Console handler to output logs to the console
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.DEBUG)

# Logging format
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
file_handler.setFormatter(formatter)
console_handler.setFormatter(formatter)

# Add handlers to the logger
logger.addHandler(file_handler)
logger.addHandler(console_handler)

# MongoDB setup
DB_NAME = os.environ.get('DB_NAME')
COLLECTION_NAME = os.environ.get('COLLECTION_NAME')
MONGO_CONNECTION_STRING = os.environ.get('MONGO_CONNECTION_STRING')

if not all([DB_NAME, COLLECTION_NAME, MONGO_CONNECTION_STRING]):
    logger.error("One or more required MongoDB environment variables are not set")
    raise ValueError("One or more required MongoDB environment variables are not set")

client = pymongo.MongoClient(MONGO_CONNECTION_STRING)
db = client[DB_NAME]
collection = db[COLLECTION_NAME]

def fetch_article_urls(base_url, pages):
    logger.info(f"Fetching article URLs from {base_url} for {pages} pages.")
    article_urls = []
    try:
        for page in range(1, pages + 1):
            url = base_url if page == 1 else f"{base_url}page/{page}/"
            response = requests.get(url)
            soup = BeautifulSoup(response.content, 'html.parser')
            for h1_tag in soup.find_all('h1', id='list'):
                a_tag = h1_tag.find('a')
                if a_tag and a_tag.get('href'):
                    article_urls.append(a_tag['href'])
    except Exception as e:
        logger.exception(f"Error while fetching article URLs: {e}")
    return article_urls

async def scrape_and_get_content(url):
    logger.info(f"Scraping content from {url}")
    content_list = []
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')
        main_content = soup.find('div', class_='inside_post column content_width')
        if not main_content:
            raise Exception("Main content div not found")

        heading = main_content.find('h1', id='list')
        if not heading:
            raise Exception("Heading not found")

        heading_text = heading.get_text()
        content_list.append({'type': 'heading', 'text': heading_text})

        for tag in main_content.find_all(recursive=False):
            if tag.get('class') in [['sharethis-inline-share-buttons', 'st-center', 'st-has-labels', 'st-inline-share-buttons', 'st-animated'], ['prenext']]:
                continue
            text = tag.get_text()
            if tag.name == 'p':
                content_list.append({'type': 'paragraph', 'text': text})
            elif tag.name == 'h2':
                content_list.append({'type': 'heading_2', 'text': text})
            elif tag.name == 'h4':
                content_list.append({'type': 'heading_4', 'text': text})
            elif tag.name == 'ul':
                for li in tag.find_all('li'):
                    li_text = li.get_text()
                    content_list.append({'type': 'list_item', 'text': f"• {li_text}"})
    except Exception as e:
        logger.exception(f"Error while scraping content from {url}: {e}")
    return content_list

def insert_content_between_placeholders(doc, content_list):
    logger.info("Inserting content between placeholders")
    start_placeholder = end_placeholder = None
    try:
        for i, para in enumerate(doc.paragraphs):
            if "START_CONTENT" in para.text:
                start_placeholder = i
            elif "END_CONTENT" in para.text:
                end_placeholder = i
                break

        if start_placeholder is None or end_placeholder is None:
            raise Exception("Could not find both placeholders")

        for i in range(end_placeholder - 1, start_placeholder, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)

        content_list = content_list[::-1]

        for content in content_list:
            if content['type'] == 'heading':
                doc.paragraphs[start_placeholder]._element.addnext(doc.add_heading(content['text'], level=1)._element)
            elif content['type'] == 'paragraph':
                doc.paragraphs[start_placeholder]._element.addnext(doc.add_paragraph(content['text'], style='Normal')._element)
            elif content['type'] == 'heading_2':
                doc.paragraphs[start_placeholder]._element.addnext(doc.add_heading(content['text'], level=2)._element)
            elif content['type'] == 'heading_4':
                doc.paragraphs[start_placeholder]._element.addnext(doc.add_heading(content['text'], level=4)._element)
            elif content['type'] == 'list_item':
                try:
                    doc.paragraphs[start_placeholder]._element.addnext(doc.add_paragraph(content['text'], style='List Bullet')._element)
                except KeyError:
                    logger.warning("Style 'List Bullet' not found, falling back to 'Normal' style for list items")
                    doc.paragraphs[start_placeholder]._element.addnext(doc.add_paragraph(content['text'], style='Normal')._element)

        doc.paragraphs[start_placeholder].text = ""
        doc.paragraphs[end_placeholder].text = ""
    except Exception as e:
        logger.exception(f"Error while inserting content: {e}")
        raise


def download_template(url):
    logger.info(f"Downloading template from {url}")
    download_url = url.replace('/edit?usp=sharing', '/export?format=docx')
    try:
        response = requests.get(download_url)
        response.raise_for_status()
        return io.BytesIO(response.content)
    except requests.exceptions.RequestException as e:
        logger.exception(f"Error while downloading template: {e}")
        raise

def check_and_insert_urls(urls):
    logger.info(f"Checking and inserting URLs")
    new_urls = []
    try:
        for url in urls:
            if 'daily-current-affairs-quiz' in url:
                continue
            if not collection.find_one({'url': url}):
                new_urls.append(url)
                collection.insert_one({'url': url})
    except Exception as e:
        logger.exception(f"Error while checking or inserting URLs: {e}")
    return new_urls

def convert_docx_to_pdf(docx_path, pdf_path):
    logger.info(f"Converting {docx_path} to PDF")
    try:
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                        os.path.dirname(pdf_path), docx_path], 
                       check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        original_pdf = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
        original_pdf_path = os.path.join(os.path.dirname(pdf_path), original_pdf)
        os.rename(original_pdf_path, pdf_path)
    except subprocess.CalledProcessError as e:
        logger.exception(f"Error while converting DOCX to PDF: {e}")
        raise

def rename_pdf(pdf_path, new_name):
    logger.info(f"Renaming PDF {pdf_path} to {new_name}")
    new_pdf_path = os.path.join(os.path.dirname(pdf_path), new_name)
    try:
        os.rename(pdf_path, new_pdf_path)
    except Exception as e:
        logger.exception(f"Error while renaming PDF: {e}")
        raise
    return new_pdf_path

async def send_pdf_to_telegram(pdf_path, bot_token, channel_id, caption):
    logger.info(f"Sending PDF to Telegram channel {channel_id}")
    bot = telegram.Bot(token=bot_token)
    for _ in range(3):
        try:
            with open(pdf_path, 'rb') as pdf_file:
                await bot.send_document(chat_id=channel_id, document=pdf_file, filename=os.path.basename(pdf_path), caption=caption)
            break
        except telegram.error.TimedOut as e:
            logger.warning(f"Timeout error while sending PDF to Telegram, retrying: {e}")
            await asyncio.sleep(5)
        except Exception as e:
            logger.exception(f"Error while sending PDF to Telegram: {e}")
            raise

async def main():
    try:
        base_url = "https://www.gktoday.in/current-affairs/"
        article_urls = fetch_article_urls(base_url, 3)
        new_urls = check_and_insert_urls(article_urls)
        if not new_urls:
            logger.info("No new URLs found to process")
            return
        
        template_url = os.environ.get('TEMPLATE_URL')
        if not template_url:
            logger.error("TEMPLATE_URL environment variable is not set")
            raise ValueError("TEMPLATE_URL environment variable is not set")
        
        template_bytes = download_template(template_url)
        
        doc = Document(template_bytes)
        
        all_content = []
        english_titles = []
        for url in new_urls:
            content_list = await scrape_and_get_content(url)
            all_content.extend(content_list)
            english_titles.append(content_list[0]['text'])  # Assuming the first item is the title
        
        insert_content_between_placeholders(doc, all_content)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
            doc.save(tmp_docx.name)
        
        pdf_path = tmp_docx.name.replace('.docx', '.pdf')
        
        convert_docx_to_pdf(tmp_docx.name, pdf_path)
        
        # Rename the PDF file
        current_date = datetime.now().strftime('%d-%m-%Y')
        new_pdf_name = f"{current_date} Current Affairs.pdf"
        renamed_pdf_path = rename_pdf(pdf_path, new_pdf_name)
        
        bot_token = os.environ.get('TELEGRAM_BOT_TOKEN')
        channel_id = os.environ.get('TELEGRAM_CHANNEL_ID')
        
        if not bot_token or not channel_id:
            logger.error("TELEGRAM_BOT_TOKEN or TELEGRAM_CHANNEL_ID environment variable is not set")
            raise ValueError("TELEGRAM_BOT_TOKEN or TELEGRAM_CHANNEL_ID environment variable is not set")
        
        caption = (
            f"🎗️ {datetime.now().strftime('%d %B %Y')} Current Affairs 🎗️\n\n"
            + '\n'.join([f"👉 {title}" for title in english_titles]) + '\n\n'
            + "🎉 Join us :- @Daily_Current_All_Source 🎉"
        )
        
        await send_pdf_to_telegram(renamed_pdf_path, bot_token, channel_id, caption)
        
        os.unlink(tmp_docx.name)
        os.unlink(renamed_pdf_path)
        
    except Exception as e:
        logger.exception(f"An error occurred in the main process: {e}")
        raise

if __name__ == "__main__":
    logger.info("Script started")
    asyncio.run(main())
    logger.info("Script finished")
